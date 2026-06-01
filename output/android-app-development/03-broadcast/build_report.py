from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/Users/cuing/AndroidStudioProjects/android-course-projects")
TEMPLATE = Path("/Users/cuing/Desktop/课内实验报告_实验3-广播.docx")
OUT_DIR = ROOT / "output/android-app-development/03-broadcast"
REPORT = OUT_DIR / "课内实验报告_实验3-广播_已完成.docx"
SCREENSHOT_DIR = OUT_DIR / "screenshots"
SOURCE_DIR = OUT_DIR / "source"


def set_run_font(run, size=12, bold=False, name="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line=False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_paragraph(doc, text="", bold=False, size=12, first_line=False, align=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text):
    p = add_paragraph(doc, text, bold=True, size=12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_info_table(doc):
    rows = [
        ("MainActivity", "负责界面初始化、动态注册广播接收者、发送无序广播、有序广播、拦截广播和指定接收者广播，并将接收顺序显示到运行结果区域。"),
        ("DuckReceiver", "继承 BroadcastReceiver，通过构造参数区分接收者名称、是否拦截广播，并在 onReceive() 中回调显示接收结果。"),
        ("StaticDuckReceiver", "在 AndroidManifest.xml 中静态声明，用于演示通过 ComponentName 指定某一个广播接收者。"),
        ("activity_main.xml", "设计实验界面，包含五个功能按钮、运行结果显示框和清空按钮。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_cell_text(table.rows[0].cells[0], "类/文件", bold=True)
    set_cell_text(table.rows[0].cells[1], "作用说明", bold=True)
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)


def add_picture(doc, image_name, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(SCREENSHOT_DIR / image_name), width=Cm(9.2))
    cap = add_paragraph(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.space_after = Pt(6)


def add_code_block(doc, title, path):
    add_heading(doc, title)
    text = path.read_text(encoding="utf-8")
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8)


def remove_placeholder_tail(doc):
    start_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("。。。"):
            start_index = idx
            break
    if start_index is None:
        return
    for paragraph in doc.paragraphs[start_index:]:
        paragraph._element.getparent().remove(paragraph._element)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_document_fonts(doc):
    for paragraph in doc.paragraphs:
        set_paragraph_format(paragraph)
        for run in paragraph.runs:
            set_run_font(run, size=12, bold=run.bold)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph)
                    for run in paragraph.runs:
                        set_run_font(run, size=12, bold=run.bold)


def main():
    doc = Document(TEMPLATE)
    set_document_fonts(doc)
    remove_placeholder_tail(doc)

    add_heading(doc, "四、程序设计和说明")
    add_paragraph(doc, "本实验新建模块 broadcastApp，包名为 cn.itcast.broadcast。程序围绕“数鸭子”案例完成五类广播操作：无序广播、有序广播（优先级不同）、有序广播（优先级相同）、拦截有序广播、指定广播接收者。界面上方提供五个按钮，点击后在下方运行结果区域显示广播发送内容和接收者接收顺序。", first_line=True)
    add_paragraph(doc, "动态广播接收者通过 registerReceiver() 注册，并通过 IntentFilter.setPriority() 设置有序广播优先级；拦截广播时，高优先级接收者在 onReceive() 中调用 abortBroadcast()；指定接收者功能使用 ComponentName 将广播发送给 StaticDuckReceiver。", first_line=True)
    add_info_table(doc)

    add_heading(doc, "五、运行结果")
    screenshots = [
        ("01_unordered.png", "图1 发送无序广播的效果"),
        ("02_ordered_priority.png", "图2 发送有序广播的效果（优先级均不同）"),
        ("03_ordered_same_priority.png", "图3 发送有序广播的效果（优先级相同）"),
        ("04_intercept.png", "图4 拦截广播的效果"),
        ("05_static_receiver.png", "图5 指定广播接收者效果"),
    ]
    for image_name, caption in screenshots:
        add_picture(doc, image_name, caption)

    add_page_break(doc)
    add_heading(doc, "六、核心源码")
    add_paragraph(doc, f"完整源码已同步整理到：{SOURCE_DIR}", first_line=True)
    source_files = [
        ("MainActivity.java", SOURCE_DIR / "MainActivity.java"),
        ("DuckReceiver.java", SOURCE_DIR / "DuckReceiver.java"),
        ("StaticDuckReceiver.java", SOURCE_DIR / "StaticDuckReceiver.java"),
        ("AndroidManifest.xml", SOURCE_DIR / "AndroidManifest.xml"),
        ("activity_main.xml", SOURCE_DIR / "activity_main.xml"),
    ]
    for title, path in source_files:
        add_code_block(doc, title, path)

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
